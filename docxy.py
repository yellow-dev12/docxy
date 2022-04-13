from docx import Document
from docx.shared import Inches
import os
import os.path
import shutil
from random import randint
from bs4 import BeautifulSoup
import requests
from icrawler.builtin import GoogleImageCrawler
import PIL

class Block:
    text = ""
    title = 0
    def __init__(self, text, title):
        self.text = text
        out = []
        if len(text.split(" ")) > 3:
            for i in range(randint(1,3)):
                out.append(text.split(" ")[i])
            self.title = " ".join(out)


#Настройки
find_text = "Антарктида" #Тема доклада
names = "Радион Агумава, Дима Иванов, Богдан, Артём, Рамин, Адам, Данил Григорчук, Даниил Цветиков".split(", ") #Имена для кого делается доклад (через запятую)
docs_count = 8 #Количество докладов
min_max_blocks_count = "5, 8" #минимальное и максимальное количество блоков для докладов (через запятую)
min_block_size = 25 #Минимальное количество слов в одном блоке
use_imgs = True #Использовать фотографии?
img_per_paragraph = 2 #Количество фоток на каждые ? параграфов (поставь ноль чтобы было рандомно)
docs_type = "docx" #Тип документа на выходе: docx для винды, odt для линукса
file_name = False #Называть документы именем человека или же просто цифрой
repeat_blocks_delete = True #Защита от повторений блоков
pics = 16 #Своё число фоток (если нужно фоток под минимум то вписать None)
rand_name = True # Рандомное имя если код не может найти нужное

#Настройки главного документа
blocks_count = 7 #Количество блоков
main_style = "1, 1" #Стиль: курсив или нет (0, 1); Положение (0 - лево, 1 - центр, 2 - право)



#Чтение с сайта
print("Создание папки")
if find_text in os.listdir(os.getcwd()): shutil.rmtree(find_text)
os.mkdir(find_text)
os.chdir(find_text)
print("Чтение с сайта")
all_blocks = []
soup = BeautifulSoup(requests.get("https://ru.wikipedia.org/wiki/" + find_text).text, "lxml")
for i in range(len(soup.find_all("p"))): all_blocks.append(Block(soup.find_all("p")[i].text, i))
print("Всего блоков найдено " + str(len(all_blocks)))
print("------------>")

#if min_max_blocks_count.split(", ")[1] > blocks_count:
#    if len(all_blocks) < min_max_blocks_count[1]:
#        print("Количество найденных блоков меньше чем нужно было найти, повторять данные блоки или создать со столькми блоками сколько удалось найти? (1, 0) по умолчанию=1")
#        inp = input(":")
#        if inp == "0":
#            min_max_blocks_count[1] = len(all_blocks)

#Фотки
if use_imgs:
    max_num = 0
    if max_num != None:
        if blocks_count > int(min_max_blocks_count.split(", ")[1]):
            max_num = blocks_count
        else:
            max_num = int(min_max_blocks_count.split(", ")[1])
    else:
        max_num = pics
    print(f"Загрузка {max_num} фотографий")
    google_Crawler = GoogleImageCrawler(storage = {'root_dir': r'images'})
    google_Crawler.crawl(keyword = find_text, max_num = max_num)
    os.chdir("images")
    for file in os.listdir():
        if ".png" in file:
            temp_img = PIL.Image.open(file)
            temp_img.convert("RGB").save(file.split(".")[0] + ".jpg")
            os.remove(file)
    os.chdir("../")


#Создание документов
#Первый документ
print("Создание первого документа...")
first_doc = Document()
first_doc.add_heading(find_text).alignment = 1
first_doc.add_heading(names[0], level=2).alignment = 1
first_doc.add_page_break()
last_pic = None
temp_blocks = []
for block in range(blocks_count):
    print("Блок номер " + str(block))
    temp_block = all_blocks[randint(0, len(all_blocks) - 1)]

    if repeat_blocks_delete:
        if temp_block in temp_blocks: break
        else: temp_blocks.append(temp_block)

    if min_block_size == len(all_blocks):
        while len(temp_block.text.split(" ")) < min_block_size: temp_block = all_blocks[randint(0, len(all_blocks) - 1)]
    else:
        while len(temp_block.text.split(" ")) < len(all_blocks): temp_block = all_blocks[randint(0, len(all_blocks) - 1)]
    if block == 0: temp_block = all_blocks[0]
    elif block == 1: temp_block = all_blocks[1]
    print("Кол-во слов: " + str(len(temp_block.text.split(" "))))
    print(temp_block.title)
    if linux: head = first_doc.add_paragraph(temp_block.title)
    else: head = first_doc.add_heading(temp_block.title, level=2)
    head.alignment = int(main_style.split(", ")[1])
    head.italic = bool(main_style.split(",")[0])
    first_doc.add_paragraph(temp_block.text)
    if use_imgs and block % img_per_paragraph == 0:
        print("Добавление фотографии")
        os.chdir("images")
        randy = randint(1, len(os.listdir(os.getcwd()))-1 )
        temp_pic = f"00000{randy}.jpg"
        first_doc.add_picture(temp_pic, width=Inches(float(randint(2, 7))))
        os.chdir("../")
    print("Блок создан")
    print("--------------------------")
print("Сохранение")
first_doc.save(f"main.{docs_type}")
if linux:
    if input("Распечатать?") == "y": os.system(f"hp-print main.{docs_type}")

#Другие документы
if docs_count > 1:
    for i in range(docs_count - 1):
        print(f"Создание {i} документа...")
        #Рандомные настройки
        rand_style = [randint(0, 1), randint(0, 2)]
        rand_blocks_count = randint(int(min_max_blocks_count.split(", ")[0]), int(min_max_blocks_count.split(", ")[1]))
        try:
            doc_name = names[i]
        except:
            if not rand_name:
                doc_name = names[len(names) - 1]
            else:
                doc_name = "Сюды имя"
        print(f"Настройки: стиль: {rand_style[0]}; положение: {rand_style[1]}; кол-во блоков: {rand_blocks_count}")
        print("^^^")
        first_doc = Document()
        first_doc.add_heading(find_text).alignment = 1
        first_doc.add_heading(names[i], level=2).alignment = 1
        first_doc.add_page_break()
        for block in range(rand_blocks_count):
            print("Блок номер " + str(block))
            temp_block = all_blocks[randint(0, len(all_blocks) - 1)]
            while len(temp_block.text.split(" ")) < min_block_size: temp_block = all_blocks[randint(0, len(all_blocks) - 1)]
            if block == 0: temp_block = all_blocks[0]
            elif block == 1: temp_block = all_blocks[1]
            print("Кол-во слов: " + str(len(temp_block.text.split(" "))))
            if linux: head = first_doc.add_paragraph(temp_block.title)
            else: head = first_doc.add_heading(temp_block.title, level=2)
            head.alignment = rand_style[1]
            head.italic = bool(rand_style[0])
            first_doc.add_paragraph(temp_block.text)
            if use_imgs and block % img_per_paragraph == 0:
                print("Добавление фотографии")
                os.chdir("images")
                randy = randint(1, len(os.listdir(os.getcwd()))-1 )
                temp_pic = f"00000{randy}.jpg"
                first_doc.add_picture(temp_pic, width=Inches(float(randint(2, 7))))
                os.chdir("../")
            print("Блок создан")
            print("--------------------------")
        print("Сохранение")
        if file_name:
            first_doc.save(f"{names[i + 1]}.{docs_type}")
            if linux:
                if input("Распечатать? ") == "y": os.system(f"{names[i + 1]}.{docs_type}")
        else:
            first_doc.save(f"{i + 1}.{docs_type}")
            if linux:
                if input("Распечатать? ") == "y": os.system(f"{i + 1}.{docs_type}")